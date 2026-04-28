"""
SQM 재고관리 시스템 - PDF/이미지 변환기 (v2.6.0)

★ PDF 및 캡처 이미지를 Excel/Word로 변환 ★

지원 입력 유형:
1. 텍스트 기반 PDF - 직접 텍스트 추출
2. 이미지/스캔 PDF - OCR로 텍스트 추출
3. 표가 있는 PDF - 테이블 구조 유지하여 Excel 변환
4. 캡처 이미지 (PNG/JPG/JPEG/BMP/TIFF) - OCR → Excel/Word 변환

사용법:
    from pdf_converter import PDFConverter
    
    converter = PDFConverter()
    
    # PDF → Excel
    excel_path = converter.to_excel("document.pdf")
    
    # 캡처 이미지 → Excel
    excel_path = converter.to_excel("capture.png")
    
    # Word로 변환 (PDF/이미지 모두 가능)
    word_path = converter.to_word("document.pdf")
    word_path = converter.to_word("screenshot.jpg")

Author: Ruby
Version: 2.6.0
"""

import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

logger = logging.getLogger(__name__)

# =============================================================================
# 의존성 체크
# =============================================================================

# PyMuPDF (PDF 처리)
try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF 미설치 - pip install pymupdf")

# openpyxl (Excel 생성)
try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    logger.warning("openpyxl 미설치 - pip install openpyxl")

# python-docx (Word 생성)
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx 미설치 - pip install python-docx")

# pytesseract (OCR) - 선택적
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    logger.info("pytesseract 미설치 - OCR 기능 제한됨")


# =============================================================================
# 데이터 클래스
# =============================================================================

@dataclass
class PageContent:
    """페이지 컨텐츠"""
    page_num: int
    text: str
    tables: List[List[List[str]]] = field(default_factory=list)  # 테이블 목록
    images: List[bytes] = field(default_factory=list)  # 이미지 목록
    is_scanned: bool = False  # 스캔 이미지 여부


@dataclass
class ConversionResult:
    """변환 결과"""
    success: bool
    source_file: str
    output_file: str
    output_type: str  # "excel" or "word"
    page_count: int
    table_count: int
    error_message: str = ""
    processing_time: float = 0.0


# =============================================================================
# PDF 변환기
# =============================================================================

IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.bmp', '.tif', '.tiff', '.webp'}


class PDFConverter:
    """
    PDF/이미지 → Excel/Word 변환기
    
    기능:
    1. 텍스트 추출 (일반 PDF)
    2. OCR 추출 (스캔/이미지 PDF)
    3. 테이블 감지 및 구조화
    4. 캡처 이미지 OCR → Excel/Word 변환
    5. Excel/Word 파일 생성
    """

    def __init__(self, output_dir: str = None, ocr_lang: str = "eng+kor"):
        """
        Args:
            output_dir: 출력 디렉토리 (기본: ./output/converted)
            ocr_lang: OCR 언어 (기본: 영어+한국어)
        """
        if output_dir is None:
            base_dir = Path(__file__).resolve().parent
            output_dir = base_dir / "output" / "converted"

        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.ocr_lang = ocr_lang

        # 의존성 체크
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """필수 라이브러리 확인"""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF가 필요합니다: pip install pymupdf")
        if not HAS_OPENPYXL:
            raise ImportError("openpyxl이 필요합니다: pip install openpyxl")

    @staticmethod
    def _is_image_file(file_path: str) -> bool:
        """이미지 파일 여부 판별"""
        return Path(file_path).suffix.lower() in IMAGE_EXTENSIONS

    # =========================================================================
    # PDF 분석
    # =========================================================================

    def analyze_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """
        PDF 파일 분석
        
        Returns:
            dict: 분석 결과 (페이지 수, 텍스트/이미지 비율, 테이블 수 등)
        """
        pdf_path = Path(pdf_path)
        if not pdf_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {pdf_path}")

        doc = None
        try:
            doc = fitz.open(str(pdf_path))

            result = {
                'filename': pdf_path.name,
                'page_count': len(doc),
                'has_text': False,
                'has_images': False,
                'is_scanned': False,
                'estimated_tables': 0,
                'text_length': 0,
                'pages': []
            }

            total_text_len = 0
            total_images = 0

            for page_num, page in enumerate(doc):
                text = page.get_text()
                images = page.get_images()

                page_info = {
                    'page': page_num + 1,
                    'text_length': len(text),
                    'image_count': len(images),
                    'has_tables': self._detect_table_structure(text)
                }

                total_text_len += len(text)
                total_images += len(images)

                if page_info['has_tables']:
                    result['estimated_tables'] += 1

                result['pages'].append(page_info)

            result['text_length'] = total_text_len
            result['has_text'] = total_text_len > 100
            result['has_images'] = total_images > 0

            # 스캔 PDF 판단: 이미지는 있는데 텍스트가 거의 없음
            if total_images > 0 and total_text_len < 50 * len(doc):
                result['is_scanned'] = True

            return result
        finally:
            if doc:
                doc.close()

    def _detect_table_structure(self, text: str) -> bool:
        """텍스트에서 테이블 구조 감지"""
        # 탭이나 연속 공백이 많으면 테이블 가능성
        if text.count('\t') > 3:
            return True
        if len(re.findall(r'  +', text)) > 5:
            return True
        # 숫자와 텍스트가 규칙적으로 반복
        if len(re.findall(r'\d+[,.]?\d*\s+\w+', text)) > 3:
            return True
        return False

    # =========================================================================
    # 텍스트 추출
    # =========================================================================

    def extract_content(self, file_path: str) -> List[PageContent]:
        """
        PDF 또는 이미지에서 모든 컨텐츠 추출.
        v2.6.0: 캡처 이미지(PNG/JPG 등) 지원 추가
        
        Args:
            file_path: PDF 또는 이미지 파일 경로
        
        Returns:
            List[PageContent]: 페이지별 컨텐츠
        """
        if self._is_image_file(file_path):
            return self._extract_content_from_image(file_path)
        return self._extract_content_from_pdf(file_path)

    def _extract_content_from_image(self, image_path: str) -> List[PageContent]:
        """캡처 이미지에서 OCR로 컨텐츠 추출 (v2.6.0)"""
        content = PageContent(page_num=1, text="", is_scanned=True)

        if not HAS_OCR:
            logger.warning("pytesseract 미설치 — 이미지 OCR 불가")
            content.text = "[OCR 불가] pytesseract가 설치되지 않았습니다."
            return [content]

        try:
            img = Image.open(image_path)
            text = pytesseract.image_to_string(img, lang=self.ocr_lang)
            content.text = text

            table = self._extract_table_from_text(text)
            if table:
                content.tables.append(table)

            logger.info(f"이미지 OCR 완료: {Path(image_path).name} ({len(text)} chars)")
        except (OSError, IOError) as e:
            logger.error(f"이미지 OCR 실패: {e}")
            content.text = f"[OCR 실패] {e}"

        return [content]

    def _extract_content_from_pdf(self, pdf_path: str) -> List[PageContent]:
        """PDF에서 모든 컨텐츠 추출 (v2.9.41: 안전한 파일 핸들링)"""
        doc = None
        try:
            doc = fitz.open(str(pdf_path))
            pages = []

            for page_num, page in enumerate(doc):
                content = PageContent(page_num=page_num + 1, text="")

                # 1. 텍스트 추출
                text = page.get_text()

                # 2. 텍스트가 없으면 OCR 시도
                if len(text.strip()) < 50 and HAS_OCR:
                    logger.info(f"페이지 {page_num + 1}: OCR 처리 중...")
                    text = self._ocr_page(page)
                    content.is_scanned = True

                content.text = text

                # 3. 테이블 추출
                tables = self._extract_tables(page)
                content.tables = tables

                pages.append(content)

            return pages
        finally:
            if doc:
                doc.close()

    def _ocr_page(self, page) -> str:
        """페이지 OCR 처리"""
        if not HAS_OCR:
            return ""

        try:
            # 페이지를 이미지로 변환
            pix = page.get_pixmap(dpi=300)
            img_data = pix.tobytes("png")

            # PIL 이미지로 변환
            from io import BytesIO
            img = Image.open(BytesIO(img_data))

            # OCR 수행
            text = pytesseract.image_to_string(img, lang=self.ocr_lang)
            return text
        except (OSError, IOError, PermissionError) as e:
            logger.warning(f"OCR 실패: {e}")
            return ""

    def _extract_tables(self, page) -> List[List[List[str]]]:
        """페이지에서 테이블 추출"""
        tables = []

        try:
            # PyMuPDF의 테이블 찾기 기능 사용
            tabs = page.find_tables()

            for tab in tabs:
                table_data = []
                for row in tab.extract():
                    # None 값을 빈 문자열로 변환
                    clean_row = [str(cell) if cell else "" for cell in row]
                    table_data.append(clean_row)

                if table_data:
                    tables.append(table_data)
        except (ValueError, TypeError, AttributeError) as e:
            logger.debug(f"테이블 추출 실패: {e}")
            # 대안: 텍스트 기반 테이블 추출 시도
            text_table = self._extract_table_from_text(page.get_text())
            if text_table:
                tables.append(text_table)

        return tables

    def _extract_table_from_text(self, text: str) -> List[List[str]]:
        """텍스트에서 테이블 구조 추출 (대안)"""
        lines = text.strip().split('\n')
        table = []

        for line in lines:
            # 탭 또는 연속 공백으로 분리
            if '\t' in line:
                cells = [c.strip() for c in line.split('\t')]
            else:
                cells = [c.strip() for c in re.split(r'  +', line)]

            # 2개 이상의 셀이 있으면 테이블 행으로 간주
            if len(cells) >= 2:
                table.append(cells)

        return table if len(table) >= 2 else []

    # =========================================================================
    # Excel 변환
    # =========================================================================

    def to_excel(
        self,
        file_path: str,
        output_path: str = None,
        include_text: bool = True,
        table_only: bool = False
    ) -> ConversionResult:
        """
        PDF 또는 캡처 이미지 → Excel 변환
        
        Args:
            file_path: PDF 또는 이미지 파일 경로
            output_path: 출력 파일 경로 (기본: 자동 생성)
            include_text: 텍스트도 포함할지 여부
            table_only: 테이블만 추출할지 여부
        
        Returns:
            ConversionResult: 변환 결과
        """
        start_time = datetime.now()
        file_path = Path(file_path)

        if output_path is None:
            output_path = self.output_dir / f"{file_path.stem}.xlsx"

        try:
            # 컨텐츠 추출 (PDF 또는 이미지)
            pages = self.extract_content(str(file_path))

            # Excel 생성
            wb = Workbook()

            # 기본 시트 제거
            if 'Sheet' in wb.sheetnames:
                del wb['Sheet']

            table_count = 0

            # 1. 테이블 시트들 생성
            for page in pages:
                for t_idx, table in enumerate(page.tables):
                    table_count += 1
                    sheet_name = f"Table_P{page.page_num}_{t_idx + 1}"
                    ws = wb.create_sheet(title=sheet_name[:31])  # Excel 시트명 31자 제한

                    self._write_table_to_sheet(ws, table)

            # 2. 텍스트 시트 생성 (옵션)
            if include_text and not table_only:
                ws_text = wb.create_sheet(title="전체 텍스트", index=0)
                self._write_text_to_sheet(ws_text, pages)

            # 시트가 하나도 없으면 기본 시트 추가
            if len(wb.sheetnames) == 0:
                ws = wb.create_sheet(title="내용")
                ws['A1'] = "추출된 테이블이 없습니다."
                for page in pages:
                    row = ws.max_row + 2
                    ws.cell(row=row, column=1, value=f"=== 페이지 {page.page_num} ===")
                    for line_idx, line in enumerate(page.text.split('\n')[:100]):
                        ws.cell(row=row + line_idx + 1, column=1, value=line)

            # 저장
            wb.save(str(output_path))

            processing_time = (datetime.now() - start_time).total_seconds()

            return ConversionResult(
                success=True,
                source_file=str(file_path),
                output_file=str(output_path),
                output_type="excel",
                page_count=len(pages),
                table_count=table_count,
                processing_time=processing_time
            )

        except (ValueError, TypeError, AttributeError) as e:
            logger.error(f"Excel 변환 실패: {e}")
            return ConversionResult(
                success=False,
                source_file=str(file_path),
                output_file=str(output_path),
                output_type="excel",
                page_count=0,
                table_count=0,
                error_message=str(e)
            )

    def _write_table_to_sheet(self, ws, table: List[List[str]]) -> None:
        """테이블을 Excel 시트에 쓰기"""
        # 스타일 정의
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        for row_idx, row in enumerate(table):
            for col_idx, cell_value in enumerate(row):
                cell = ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell_value)
                cell.border = border
                cell.alignment = Alignment(wrap_text=True, vertical='top')

                # 첫 행은 헤더 스타일
                if row_idx == 0:
                    cell.fill = header_fill
                    cell.font = header_font

        # 열 너비 자동 조정
        for col_idx in range(1, ws.max_column + 1):
            max_length = 0
            for row_idx in range(1, ws.max_row + 1):
                cell_value = ws.cell(row=row_idx, column=col_idx).value
                if cell_value:
                    max_length = max(max_length, len(str(cell_value)))

            adjusted_width = min(max(max_length + 2, 10), 50)
            ws.column_dimensions[get_column_letter(col_idx)].width = adjusted_width

    def _write_text_to_sheet(self, ws, pages: List[PageContent]) -> None:
        """텍스트를 Excel 시트에 쓰기"""
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 100

        ws['A1'] = "페이지"
        ws['B1'] = "내용"
        ws['A1'].font = Font(bold=True)
        ws['B1'].font = Font(bold=True)

        row = 2
        for page in pages:
            ws.cell(row=row, column=1, value=f"페이지 {page.page_num}")

            # 텍스트를 줄바꿈 포함하여 저장
            text_cell = ws.cell(row=row, column=2, value=page.text[:32000])  # Excel 셀 제한
            text_cell.alignment = Alignment(wrap_text=True, vertical='top')

            row += 1

    # =========================================================================
    # Word 변환
    # =========================================================================

    def to_word(
        self,
        file_path: str,
        output_path: str = None,
        preserve_layout: bool = True
    ) -> ConversionResult:
        """
        PDF 또는 캡처 이미지 → Word 변환
        
        Args:
            file_path: PDF 또는 이미지 파일 경로
            output_path: 출력 파일 경로 (기본: 자동 생성)
            preserve_layout: 레이아웃 유지 시도 여부
        
        Returns:
            ConversionResult: 변환 결과
        """
        if not HAS_DOCX:
            return ConversionResult(
                success=False,
                source_file=str(file_path),
                output_file="",
                output_type="word",
                page_count=0,
                table_count=0,
                error_message="python-docx가 필요합니다: pip install python-docx"
            )

        start_time = datetime.now()
        file_path = Path(file_path)

        if output_path is None:
            output_path = self.output_dir / f"{file_path.stem}.docx"

        try:
            # 컨텐츠 추출 (PDF 또는 이미지)
            pages = self.extract_content(str(file_path))

            # Word 문서 생성
            doc = Document()

            # 제목
            is_image = self._is_image_file(str(file_path))
            title_prefix = "이미지 변환" if is_image else "PDF 변환"
            doc.add_heading(f"{title_prefix}: {file_path.name}", level=0)
            doc.add_paragraph(f"변환 일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")

            table_count = 0

            for page in pages:
                # 페이지 구분선
                doc.add_heading(f"페이지 {page.page_num}", level=1)

                # 테이블 추가
                for t_idx, table in enumerate(page.tables):
                    table_count += 1
                    doc.add_heading(f"테이블 {t_idx + 1}", level=2)

                    if table:
                        self._add_table_to_doc(doc, table)
                        doc.add_paragraph("")

                # 텍스트 추가
                if page.text.strip():
                    if page.tables:
                        doc.add_heading("기타 텍스트", level=2)

                    # 텍스트를 단락으로 분리
                    paragraphs = page.text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())

                # 페이지 나누기 (마지막 페이지 제외)
                if page.page_num < len(pages):
                    doc.add_page_break()

            # 저장
            doc.save(str(output_path))

            processing_time = (datetime.now() - start_time).total_seconds()

            return ConversionResult(
                success=True,
                source_file=str(file_path),
                output_file=str(output_path),
                output_type="word",
                page_count=len(pages),
                table_count=table_count,
                processing_time=processing_time
            )

        except (ValueError, TypeError, AttributeError) as e:
            logger.error(f"Word 변환 실패: {e}")
            return ConversionResult(
                success=False,
                source_file=str(file_path),
                output_file=str(output_path),
                output_type="word",
                page_count=0,
                table_count=0,
                error_message=str(e)
            )

    def _add_table_to_doc(self, doc, table: List[List[str]]) -> None:
        """Word 문서에 테이블 추가"""
        if not table:
            return

        # 최대 열 수 계산
        max_cols = max(len(row) for row in table)

        # 테이블 생성
        word_table = doc.add_table(rows=len(table), cols=max_cols)
        word_table.style = 'Table Grid'

        for row_idx, row in enumerate(table):
            for col_idx, cell_value in enumerate(row):
                if col_idx < max_cols:
                    cell = word_table.cell(row_idx, col_idx)
                    cell.text = str(cell_value) if cell_value else ""

                    # 첫 행 굵게
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    # =========================================================================
    # 통합 변환
    # =========================================================================

    def convert_all(
        self,
        pdf_path: str,
        formats: List[str] = None
    ) -> Dict[str, ConversionResult]:
        """
        PDF를 여러 형식으로 변환
        
        Args:
            pdf_path: PDF 파일 경로
            formats: 변환할 형식 목록 (기본: ["excel", "word"])
        
        Returns:
            Dict[str, ConversionResult]: 형식별 변환 결과
        """
        if formats is None:
            formats = ["excel", "word"]

        results = {}

        if "excel" in formats:
            results["excel"] = self.to_excel(pdf_path)

        if "word" in formats:
            results["word"] = self.to_word(pdf_path)

        return results

    def batch_convert(
        self,
        pdf_folder: str,
        output_format: str = "excel"
    ) -> List[ConversionResult]:
        """
        폴더 내 모든 PDF 일괄 변환
        
        Args:
            pdf_folder: PDF 폴더 경로
            output_format: 출력 형식 ("excel" 또는 "word")
        
        Returns:
            List[ConversionResult]: 변환 결과 목록
        """
        pdf_folder = Path(pdf_folder)
        results = []

        pdf_files = list(pdf_folder.glob("*.pdf")) + list(pdf_folder.glob("*.PDF"))

        for pdf_file in pdf_files:
            logger.info(f"변환 중: {pdf_file.name}")

            if output_format == "excel":
                result = self.to_excel(str(pdf_file))
            else:
                result = self.to_word(str(pdf_file))

            results.append(result)

        return results


# =============================================================================
# 편의 함수
# =============================================================================

def pdf_to_excel(pdf_path: str, output_path: str = None) -> str:
    """
    PDF → Excel 간편 변환
    
    Args:
        pdf_path: PDF 파일 경로
        output_path: 출력 경로 (옵션)
    
    Returns:
        str: 생성된 Excel 파일 경로
    """
    converter = PDFConverter()
    result = converter.to_excel(pdf_path, output_path)

    if result.success:
        return result.output_file
    else:
        raise Exception(f"변환 실패: {result.error_message}")


def pdf_to_word(pdf_path: str, output_path: str = None) -> str:
    """
    PDF → Word 간편 변환
    
    Args:
        pdf_path: PDF 파일 경로
        output_path: 출력 경로 (옵션)
    
    Returns:
        str: 생성된 Word 파일 경로
    """
    converter = PDFConverter()
    result = converter.to_word(pdf_path, output_path)

    if result.success:
        return result.output_file
    else:
        raise Exception(f"변환 실패: {result.error_message}")


# =============================================================================
# 테스트
# =============================================================================

if __name__ == "__main__":
    logger.debug("=" * 60)
    logger.debug("PDF 변환기 테스트")
    logger.debug("=" * 60)

    # 의존성 확인
    logger.debug("\n[의존성 확인]")
    logger.debug(f"  PyMuPDF: {'✅' if HAS_PYMUPDF else '❌'}")
    logger.debug(f"  openpyxl: {'✅' if HAS_OPENPYXL else '❌'}")
    logger.debug(f"  python-docx: {'✅' if HAS_DOCX else '❌'}")
    logger.debug(f"  pytesseract (OCR): {'✅' if HAS_OCR else '⚠️ (선택적)'}")

    if HAS_PYMUPDF and HAS_OPENPYXL:
        logger.debug("\n✅ PDF → Excel 변환 가능")

    if HAS_PYMUPDF and HAS_DOCX:
        logger.debug("✅ PDF → Word 변환 가능")

    logger.debug("\n[사용법]")
    logger.debug("  from utils.pdf_converter import PDFConverter, pdf_to_excel, pdf_to_word")
    logger.debug("  ")
    logger.debug("  # 방법 1: 간편 함수")
    logger.debug("  excel_path = pdf_to_excel('document.pdf')")
    logger.debug("  word_path = pdf_to_word('document.pdf')")
    logger.debug("  ")
    logger.debug("  # 방법 2: 컨버터 클래스")
    logger.debug("  converter = PDFConverter()")
    logger.debug("  result = converter.to_excel('document.pdf')")
    logger.debug("  print(f'변환 완료: {result.output_file}')")
